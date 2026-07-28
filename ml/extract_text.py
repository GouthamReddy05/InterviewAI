import io
import os
from typing import Optional

import docx

try:
    # pypdf is the maintained successor to PyPDF2; fall back if it is absent.
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - depends on which package is installed
    from PyPDF2 import PdfReader


def extract_text_from_file(file_source, filename: Optional[str] = None) -> str:
    """Extract plain text from a PDF or DOCX.

    Supports a FastAPI ``UploadFile``, any file-like object, or a path. When the
    source carries no name of its own (e.g. a bare ``BytesIO``), pass ``filename``
    so the format can be determined.
    """
    text = ""

    if hasattr(file_source, "file") and hasattr(file_source, "filename"):
        # FastAPI UploadFile
        resolved_name = filename or file_source.filename or ""
        content_stream = io.BytesIO(file_source.file.read())
        file_source.file.seek(0)

    elif hasattr(file_source, "read"):
        resolved_name = filename or getattr(file_source, "name", "") or getattr(
            file_source, "filename", ""
        )
        content_stream = io.BytesIO(file_source.read())
        if hasattr(file_source, "seek"):
            file_source.seek(0)

    elif isinstance(file_source, str):
        if not os.path.exists(file_source):
            raise FileNotFoundError(f"No such file: '{file_source}'")
        resolved_name = filename or file_source
        content_stream = file_source

    else:
        raise TypeError("Invalid file source provided.")

    resolved_name = str(resolved_name).lower()

    if resolved_name.endswith(".pdf"):
        reader = PdfReader(content_stream)
        # An encrypted PDF raises an opaque error deep inside the parser; give a
        # message the caller can surface to a user instead.
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("This PDF is password protected.") from exc

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    elif resolved_name.endswith(".docx"):
        document = docx.Document(content_stream)
        for para in document.paragraphs:
            text += para.text + "\n"
        # Resumes frequently keep contact details and skills inside tables, which
        # the paragraph walk above skips entirely.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text += " | ".join(cells) + "\n"

    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")

    return text.strip()
